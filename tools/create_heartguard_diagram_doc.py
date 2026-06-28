from pathlib import Path
import math
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas as pdf_canvas


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables" / "heartguard_chapter3_diagrams"
IMG_DIR = OUT_DIR / "images"
DOCX_PATH = OUT_DIR / "HeartGuard_Chapter_3_Design_Diagrams.docx"
PDF_PATH = OUT_DIR / "HeartGuard_Chapter_3_Design_Diagrams.pdf"


COLORS = {
    "white": "#FFFFFF",
    "ink": "#111827",
    "muted": "#475569",
    "blue": "#2563EB",
    "blue_fill": "#EFF6FF",
    "green": "#16A34A",
    "green_fill": "#ECFDF5",
    "teal": "#0F766E",
    "teal_fill": "#F0FDFA",
    "orange": "#EA580C",
    "orange_fill": "#FFF7ED",
    "red": "#DC2626",
    "red_fill": "#FEF2F2",
    "pink": "#DB2777",
    "pink_fill": "#FDF2F8",
    "purple": "#7C3AED",
    "purple_fill": "#F5F3FF",
    "slate": "#475569",
    "slate_fill": "#F8FAFC",
    "grid": "#CBD5E1",
}


def font(size, bold=False):
    candidates = [
        "arialbd.ttf" if bold else "arial.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            pass
    return ImageFont.load_default()


TITLE_FONT = font(42, True)
SUBTITLE_FONT = font(24, False)
BOX_TITLE_FONT = font(27, True)
BOX_FONT = font(22, False)
SMALL_FONT = font(18, False)
CAPTION_FONT = font(20, True)


def text_size(draw, text, fnt):
    box = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=5, align="center")
    return box[2] - box[0], box[3] - box[1]


def wrapped(text, width_chars):
    lines = []
    for raw in text.split("\n"):
        lines.extend(textwrap.wrap(raw, width=width_chars) or [""])
    return "\n".join(lines)


def rounded_box(draw, xy, title, body="", fill="#FFFFFF", outline="#2563EB", radius=28, width=4):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)
    cx = (x1 + x2) / 2
    if body:
        title_text = wrapped(title, 24)
        body_text = wrapped(body, 30)
        tw, th = text_size(draw, title_text, BOX_TITLE_FONT)
        bw, bh = text_size(draw, body_text, BOX_FONT)
        total = th + 12 + bh
        draw.multiline_text((cx, y1 + (y2 - y1 - total) / 2), title_text, anchor="ma", font=BOX_TITLE_FONT, fill=COLORS["ink"], align="center", spacing=5)
        draw.multiline_text((cx, y1 + (y2 - y1 - total) / 2 + th + 14), body_text, anchor="ma", font=BOX_FONT, fill=COLORS["muted"], align="center", spacing=5)
    else:
        title_text = wrapped(title, 26)
        draw.multiline_text((cx, (y1 + y2) / 2), title_text, anchor="mm", font=BOX_TITLE_FONT, fill=COLORS["ink"], align="center", spacing=5)


def arrow(draw, start, end, color="#2563EB", width=5):
    draw.line([start, end], fill=color, width=width)
    ang = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 18
    p1 = (end[0] - size * math.cos(ang - math.pi / 7), end[1] - size * math.sin(ang - math.pi / 7))
    p2 = (end[0] - size * math.cos(ang + math.pi / 7), end[1] - size * math.sin(ang + math.pi / 7))
    draw.polygon([end, p1, p2], fill=color)


def elbow_arrow(draw, points, color="#2563EB", width=5):
    if len(points) < 2:
        return
    for idx in range(len(points) - 2):
        draw.line([points[idx], points[idx + 1]], fill=color, width=width)
    arrow(draw, points[-2], points[-1], color=color, width=width)


def canvas(title, subtitle):
    img = Image.new("RGB", (1800, 1100), COLORS["white"])
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1800, 1100), fill=COLORS["white"])
    draw.text((90, 62), title, font=TITLE_FONT, fill=COLORS["ink"])
    draw.text((90, 112), subtitle, font=SUBTITLE_FONT, fill=COLORS["muted"])
    draw.line((90, 158, 1710, 158), fill=COLORS["grid"], width=3)
    return img, draw


def save(img, name):
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    path = IMG_DIR / name
    img.save(path, "PNG", optimize=True)
    return path


def architecture_diagram():
    img, draw = canvas("Figure 3.1: System Architecture", "Layered architecture for HeartGuard")
    boxes = [
        ((145, 250, 445, 390), "User Layer", "Doctor / Administrator", COLORS["blue_fill"], COLORS["blue"]),
        ((585, 250, 915, 390), "Frontend UI", "Forms, dashboard, reports", COLORS["blue_fill"], COLORS["blue"]),
        ((1055, 250, 1415, 390), "FastAPI Backend", "Validation, APIs, business logic", COLORS["teal_fill"], COLORS["teal"]),
        ((190, 565, 520, 725), "Authentication", "Login and role validation", COLORS["purple_fill"], COLORS["purple"]),
        ((615, 565, 975, 725), "Prediction Engine", "XGBoost readmission model", COLORS["green_fill"], COLORS["green"]),
        ((1070, 565, 1430, 725), "SQL Server Database", "Users, patients, records, reports", COLORS["orange_fill"], COLORS["orange"]),
        ((615, 850, 975, 1010), "Risk Prediction Output", "Risk score, category, confidence", COLORS["red_fill"], COLORS["red"]),
        ((1070, 850, 1430, 1010), "Power BI Dashboard", "KPIs, trends, risk distribution", COLORS["pink_fill"], COLORS["pink"]),
    ]
    for xy, title, body, fill, outline in boxes:
        rounded_box(draw, xy, title, body, fill, outline)
    arrow(draw, (445, 320), (585, 320))
    arrow(draw, (915, 320), (1055, 320))
    elbow_arrow(draw, [(1235, 390), (1235, 475), (355, 475), (355, 565)], COLORS["teal"])
    elbow_arrow(draw, [(1235, 390), (1235, 475), (795, 475), (795, 565)], COLORS["teal"])
    arrow(draw, (1235, 390), (1250, 565), COLORS["teal"])
    arrow(draw, (795, 725), (795, 850), COLORS["green"])
    arrow(draw, (1250, 725), (1250, 850), COLORS["orange"])
    return save(img, "figure_3_1_system_architecture.png")


def use_case_diagram():
    img, draw = canvas("Figure 3.2: Use Case Diagram", "Actors and major system interactions")
    draw.rounded_rectangle((405, 220, 1535, 1010), radius=34, fill="#FFFFFF", outline=COLORS["blue"], width=5)
    draw.text((735, 242), "HeartGuard System Boundary", font=BOX_TITLE_FONT, fill=COLORS["blue"])
    actors = [
        ((105, 395, 345, 515), "Doctor", COLORS["blue_fill"], COLORS["blue"]),
        ((105, 705, 345, 825), "Administrator", COLORS["green_fill"], COLORS["green"]),
        ((1555, 720, 1765, 840), "Patient\n(Optional)", COLORS["orange_fill"], COLORS["orange"]),
    ]
    for xy, label, fill, outline in actors:
        rounded_box(draw, xy, label, "", fill, outline)

    groups = [
        ((500, 325, 870, 910), "Doctor Use Cases", COLORS["blue"], COLORS["blue_fill"], [
            ("Login", COLORS["slate_fill"], COLORS["slate"]),
            ("Add / Update Patient", COLORS["blue_fill"], COLORS["blue"]),
            ("Predict Risk", COLORS["green_fill"], COLORS["green"]),
            ("View Dashboard", COLORS["pink_fill"], COLORS["pink"]),
            ("View / Download Report", COLORS["purple_fill"], COLORS["purple"]),
        ]),
        ((1015, 325, 1385, 745), "Administrator Use Cases", COLORS["green"], COLORS["green_fill"], [
            ("Login", COLORS["slate_fill"], COLORS["slate"]),
            ("Manage Users", COLORS["orange_fill"], COLORS["orange"]),
            ("Monitor System", COLORS["teal_fill"], COLORS["teal"]),
        ]),
        ((1015, 795, 1385, 925), "Patient Use Case", COLORS["orange"], COLORS["orange_fill"], [
            ("View Personal Report", COLORS["red_fill"], COLORS["red"]),
        ]),
    ]

    for xy, heading, outline, fill, cases in groups:
        x1, y1, x2, y2 = xy
        draw.rounded_rectangle(xy, radius=26, fill="#FFFFFF", outline=outline, width=4)
        draw.rounded_rectangle((x1, y1, x2, y1 + 68), radius=26, fill=fill, outline=outline, width=4)
        draw.text(((x1 + x2) / 2, y1 + 38), heading, font=BOX_TITLE_FONT, fill=COLORS["ink"], anchor="mm")
        case_y = y1 + 105
        for label, case_fill, case_outline in cases:
            rounded_box(draw, (x1 + 45, case_y, x2 - 45, case_y + 75), label, "", case_fill, case_outline, radius=38, width=3)
            case_y += 92

    arrow(draw, (345, 455), (500, 455), COLORS["blue"], 4)
    arrow(draw, (345, 765), (1015, 535), COLORS["green"], 4)
    arrow(draw, (1555, 780), (1385, 860), COLORS["orange"], 4)
    return save(img, "figure_3_2_use_case.png")


def data_flow_diagram():
    img, draw = canvas("Figure 3.3: Data Flow Diagram", "Movement of data from input to clinical insight")
    boxes = [
        ((120, 280, 430, 430), "Patient Health Data", "Age, BP, BMI, glucose, cholesterol", COLORS["blue_fill"], COLORS["blue"]),
        ((560, 280, 870, 430), "Validation", "Missing and invalid values", COLORS["teal_fill"], COLORS["teal"]),
        ((1000, 280, 1310, 430), "Preprocessing", "Encoding, scaling, cleaning", COLORS["slate_fill"], COLORS["slate"]),
        ((1440, 280, 1710, 430), "XGBoost Model", "Readmission classifier", COLORS["green_fill"], COLORS["green"]),
        ((1000, 620, 1310, 770), "Risk Classification", "Low, medium, high", COLORS["red_fill"], COLORS["red"]),
        ((560, 620, 870, 770), "Risk Score", "Probability and confidence", COLORS["orange_fill"], COLORS["orange"]),
        ((145, 825, 430, 975), "Doctor Dashboard", "Visual reports", COLORS["pink_fill"], COLORS["pink"]),
        ((560, 825, 870, 975), "PDF Report", "Downloadable summary", COLORS["purple_fill"], COLORS["purple"]),
        ((1000, 825, 1310, 975), "SQL Server", "Stored prediction record", COLORS["orange_fill"], COLORS["orange"]),
    ]
    for xy, title, body, fill, outline in boxes:
        rounded_box(draw, xy, title, body, fill, outline)
    arrow(draw, (430, 355), (560, 355))
    arrow(draw, (870, 355), (1000, 355), COLORS["teal"])
    arrow(draw, (1310, 355), (1440, 355), COLORS["green"])
    arrow(draw, (1575, 430), (1155, 620), COLORS["green"])
    arrow(draw, (1000, 695), (870, 695), COLORS["red"])
    arrow(draw, (715, 770), (287, 825), COLORS["orange"])
    arrow(draw, (715, 770), (715, 825), COLORS["orange"])
    arrow(draw, (1000, 770), (1155, 825), COLORS["red"])
    return save(img, "figure_3_3_data_flow.png")


def database_diagram():
    img, draw = canvas("Figure 3.4: Entity Relationship Diagram", "Database entities used by HeartGuard")
    entities = [
        ((105, 260, 465, 535), "USERS", ["user_id PK", "username", "password_hash", "role", "created_at"], COLORS["blue_fill"], COLORS["blue"]),
        ((720, 235, 1080, 540), "PATIENTS", ["patient_id PK", "name", "age", "gender", "contact"], COLORS["green_fill"], COLORS["green"]),
        ((1335, 260, 1695, 560), "HEALTH_RECORDS", ["record_id PK", "patient_id FK", "bmi", "cholesterol", "glucose", "blood_pressure"], COLORS["teal_fill"], COLORS["teal"]),
        ((495, 715, 855, 1010), "PREDICTIONS", ["prediction_id PK", "patient_id FK", "record_id FK", "risk_score", "risk_category"], COLORS["red_fill"], COLORS["red"]),
        ((1080, 715, 1440, 1010), "REPORTS", ["report_id PK", "patient_id FK", "prediction_id FK", "generated_by FK", "report_path"], COLORS["purple_fill"], COLORS["purple"]),
    ]
    centers = {}
    for xy, title, fields, fill, outline in entities:
        x1, y1, x2, y2 = xy
        centers[title] = ((x1 + x2) / 2, (y1 + y2) / 2)
        draw.rounded_rectangle(xy, radius=22, fill=fill, outline=outline, width=4)
        draw.rectangle((x1, y1, x2, y1 + 58), fill=outline)
        draw.text(((x1 + x2) / 2, y1 + 30), title, font=BOX_TITLE_FONT, fill="#FFFFFF", anchor="mm")
        y = y1 + 82
        for field in fields:
            draw.text((x1 + 30, y), field, font=BOX_FONT, fill=COLORS["ink"])
            y += 38
    arrow(draw, (1080, 385), (1335, 395), COLORS["teal"], 4)
    draw.text((1135, 338), "has", font=SMALL_FONT, fill=COLORS["muted"])
    arrow(draw, (900, 540), (675, 715), COLORS["red"], 4)
    draw.text((690, 610), "receives", font=SMALL_FONT, fill=COLORS["muted"])
    arrow(draw, (1515, 560), (740, 715), COLORS["teal"], 4)
    draw.text((1135, 650), "used for", font=SMALL_FONT, fill=COLORS["muted"])
    arrow(draw, (855, 860), (1080, 860), COLORS["purple"], 4)
    draw.text((925, 815), "included in", font=SMALL_FONT, fill=COLORS["muted"])
    arrow(draw, (465, 430), (1080, 790), COLORS["blue"], 4)
    draw.text((515, 520), "generates", font=SMALL_FONT, fill=COLORS["muted"])
    arrow(draw, (900, 540), (1260, 715), COLORS["green"], 4)
    draw.text((1040, 622), "owns", font=SMALL_FONT, fill=COLORS["muted"])
    return save(img, "figure_3_4_database_er.png")


def workflow_diagram():
    img, draw = canvas("Figure 3.5: Prediction Workflow", "Operational workflow for generating a readmission forecast")
    steps = [
        ((150, 245, 440, 365), "Start", "", COLORS["white"], COLORS["blue"]),
        ((585, 245, 875, 365), "User Login", "", COLORS["blue_fill"], COLORS["blue"]),
        ((1020, 245, 1310, 365), "Enter Patient Details", "", COLORS["teal_fill"], COLORS["teal"]),
        ((1455, 245, 1710, 365), "Validate Input", "", COLORS["slate_fill"], COLORS["slate"]),
        ((1180, 500, 1510, 635), "Valid Data?", "If no, show error and return to form", COLORS["orange_fill"], COLORS["orange"]),
        ((720, 500, 1010, 635), "FastAPI Backend", "Receive request", COLORS["blue_fill"], COLORS["blue"]),
        ((285, 500, 575, 635), "Preprocessing", "Apply saved transforms", COLORS["teal_fill"], COLORS["teal"]),
        ((285, 760, 575, 895), "XGBoost Model", "Generate prediction", COLORS["green_fill"], COLORS["green"]),
        ((720, 760, 1010, 895), "Display Result", "Risk score and category", COLORS["pink_fill"], COLORS["pink"]),
        ((1180, 760, 1510, 895), "Store and Report", "Database, PDF, dashboard", COLORS["purple_fill"], COLORS["purple"]),
        ((1560, 760, 1710, 895), "End", "", COLORS["white"], COLORS["blue"]),
    ]
    for xy, title, body, fill, outline in steps:
        rounded_box(draw, xy, title, body, fill, outline)
    arrow(draw, (440, 305), (585, 305))
    arrow(draw, (875, 305), (1020, 305))
    arrow(draw, (1310, 305), (1455, 305))
    arrow(draw, (1582, 365), (1438, 500), COLORS["orange"])
    arrow(draw, (1180, 568), (1010, 568), COLORS["blue"])
    arrow(draw, (720, 568), (575, 568), COLORS["teal"])
    arrow(draw, (430, 635), (430, 760), COLORS["green"])
    arrow(draw, (575, 828), (720, 828), COLORS["green"])
    arrow(draw, (1010, 828), (1180, 828), COLORS["purple"])
    arrow(draw, (1510, 828), (1560, 828), COLORS["purple"])
    arrow(draw, (1345, 500), (1165, 365), COLORS["red"], 4)
    draw.text((1172, 402), "Invalid", font=SMALL_FONT, fill=COLORS["red"])
    draw.text((1042, 537), "Valid", font=SMALL_FONT, fill=COLORS["blue"])
    return save(img, "figure_3_5_prediction_workflow.png")


def make_docx(image_paths):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HeartGuard Chapter 3 Design Diagrams")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(17, 24, 39)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("HEARTGUARD: Preventive Insights for Heart Failure Readmission Forecasting using Machine Learning")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(71, 85, 105)

    intro = doc.add_paragraph(
        "The following figures are prepared for Chapter 3: Analysis and Design. "
        "Each diagram uses a white background and a formal academic visual style suitable for insertion into the final BE Computer Engineering project report."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    placements = [
        ("Place after Section 3.6: System Architecture Design", "Figure 3.1: System Architecture of HeartGuard"),
        ("Place after Section 3.7: Use Case Analysis", "Figure 3.2: Use Case Diagram of HeartGuard"),
        ("Place after Section 3.8: Data Flow Analysis", "Figure 3.3: Data Flow Diagram of HeartGuard"),
        ("Place after Section 3.9: Database Design", "Figure 3.4: Entity Relationship Diagram of HeartGuard Database"),
        ("Place after Section 3.10 or 3.11: Prediction Workflow", "Figure 3.5: HeartGuard Prediction Workflow"),
    ]

    for idx, (path, (placement, caption)) in enumerate(zip(image_paths, placements)):
        if idx:
            doc.add_page_break()
        p = doc.add_paragraph()
        run = p.add_run(placement)
        run.bold = True
        run.font.color.rgb = RGBColor(46, 116, 181)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(8)

        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(path), width=Inches(6.5))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(5)
        cap.paragraph_format.space_after = Pt(12)
        r = cap.add_run(caption)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(71, 85, 105)

    doc.save(DOCX_PATH)
    return DOCX_PATH


def make_pdf(image_paths):
    placements = [
        ("Place after Section 3.6: System Architecture Design", "Figure 3.1: System Architecture of HeartGuard"),
        ("Place after Section 3.7: Use Case Analysis", "Figure 3.2: Use Case Diagram of HeartGuard"),
        ("Place after Section 3.8: Data Flow Analysis", "Figure 3.3: Data Flow Diagram of HeartGuard"),
        ("Place after Section 3.9: Database Design", "Figure 3.4: Entity Relationship Diagram of HeartGuard Database"),
        ("Place after Section 3.10 or 3.11: Prediction Workflow", "Figure 3.5: HeartGuard Prediction Workflow"),
    ]
    page_w, page_h = landscape(letter)
    c = pdf_canvas.Canvas(str(PDF_PATH), pagesize=(page_w, page_h))

    c.setFillColor(colors.white)
    c.rect(0, 0, page_w, page_h, stroke=0, fill=1)
    c.setFillColor(colors.HexColor("#111827"))
    c.setFont("Helvetica-Bold", 22)
    c.drawCentredString(page_w / 2, page_h - 1.5 * inch, "HeartGuard Chapter 3 Design Diagrams")
    c.setFillColor(colors.HexColor("#475569"))
    c.setFont("Helvetica", 11)
    c.drawCentredString(
        page_w / 2,
        page_h - 1.85 * inch,
        "HEARTGUARD: Preventive Insights for Heart Failure Readmission Forecasting using Machine Learning",
    )
    c.setFont("Helvetica", 10)
    c.drawCentredString(page_w / 2, page_h - 2.25 * inch, "White-background figures prepared for Chapter 3: Analysis and Design")
    c.showPage()

    for path, (placement, caption) in zip(image_paths, placements):
        c.setFillColor(colors.white)
        c.rect(0, 0, page_w, page_h, stroke=0, fill=1)
        c.setFillColor(colors.HexColor("#2E74B5"))
        c.setFont("Helvetica-Bold", 12)
        c.drawString(0.55 * inch, page_h - 0.55 * inch, placement)

        img_w = 10.0 * inch
        img_h = img_w * (1100 / 1800)
        x = (page_w - img_w) / 2
        y = 1.15 * inch
        c.drawImage(str(path), x, y, width=img_w, height=img_h, preserveAspectRatio=True, mask="auto")

        c.setFillColor(colors.HexColor("#475569"))
        c.setFont("Helvetica-Bold", 10)
        c.drawCentredString(page_w / 2, 0.65 * inch, caption)
        c.showPage()

    c.save()
    return PDF_PATH


def main():
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    images = [
        architecture_diagram(),
        use_case_diagram(),
        data_flow_diagram(),
        database_diagram(),
        workflow_diagram(),
    ]
    docx_path = make_docx(images)
    pdf_path = make_pdf(images)
    print(docx_path)
    print(pdf_path)


if __name__ == "__main__":
    main()
